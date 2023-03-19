import os
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches

# Set the file paths
pdf_file_path = 'pdf to odt\\file.pdf'
logo_file_path = 'pdf to odt\\logo.png'

# Check if the PDF file exists
if not os.path.exists(pdf_file_path):
    print(f'Error: File "{pdf_file_path}" not found')
    exit()

# Open the PDF file and extract text and logo
with open(pdf_file_path, 'rb') as pdf_file:
    pdf_reader = PdfReader(pdf_file)
    page = pdf_reader.pages[0]
    text = page.extract_text()
    # Extract logo here

# Create a new OpenOffice document
document = Document()

# Add logo
# document.add_picture(logo_file_path, width=Inches(1.25))

# Add text
document.add_paragraph(text)

# Save the document as an OpenOffice file
document.save('pdf to odt\\bill.odt')
